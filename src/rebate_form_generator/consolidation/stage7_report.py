"""Stage 7 — Generate per-supplier Word contract reports.

For each selected supplier:
  1. Reads the latest Excel from *supplier_info_dir* for field values.
  2. Opens a fresh copy of the latest Word template from *template_dir*.
  3. Replaces placeholder keywords in the document.
  4. Fills the Product Rebate Table from the corresponding
     ``contract input - <Supplier>.xlsx`` in *rebate_form_input_dir*.
  5. Saves to *output_dir/<YYYYMMDD>/Rebate Agreement Update Form#<V>_<Supplier>.docx*.

Keyword map (Word placeholder → supplier info column):
  <Supplier Name>   → Supplier Name
  <Contract Number> → Contract Number
  <Version>         → user-supplied form number
  <Name of Entity>  → Name of Entity
  <Address>         → Address
  <Signer>          → Signer
  <Title>           → Title
  <SUPPLIER-SIGN>   → SUPPLIER-Sign
"""
from __future__ import annotations

import re
from datetime import date, datetime
from pathlib import Path
from typing import Callable

from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

# Detect keyword placeholders like <Supplier Name>, <SUPPLIER-SIGN>, etc.
_KEYWORD_RE = re.compile(r"<[A-Za-z]")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_latest_file(folder: Path, extension: str) -> Path | None:
    """Return the most recently modified file with *extension* (no Office temp files)."""
    files = [f for f in folder.glob(f"*{extension}") if not f.name.startswith("~$")]
    return max(files, key=lambda f: f.stat().st_mtime) if files else None


def _normalize(text: str) -> str:
    """Collapse whitespace and lower-case for header comparison."""
    return re.sub(r"\s+", " ", str(text or "")).strip().lower()


def _str(value) -> str:
    """Safe string conversion; empty string for None."""
    return "" if value is None else str(value).strip()


def _format_cell(value, header: str = "") -> str:
    """Format a cell value for display in the Word table."""
    if value is None:
        return ""
    if isinstance(value, (date, datetime)):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, (int, float)):
        h = header.lower()
        if "rebate amount" in h or "usd" in h:
            return f"${value:,.2f}"
        if "size" in h:
            # Show as integer if value is a whole number
            return str(int(value)) if value == int(value) else str(value)
        return f"{value:,.2f}"
    return str(value)


# ---------------------------------------------------------------------------
# Supplier info
# ---------------------------------------------------------------------------

def _read_supplier_info(info_path: Path) -> dict[str, dict]:
    """Return ``{supplier_name: {column: value}}`` from the supplier info Excel."""
    wb = load_workbook(info_path, data_only=True)
    ws = wb.active
    headers = [_str(cell.value) for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    result: dict[str, dict] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        row_data = {h: v for h, v in zip(headers, row)}
        supplier_name = _str(row_data.get("Supplier Name"))
        if supplier_name:
            result[supplier_name] = row_data
    wb.close()
    return result


def _lookup_supplier(supplier_info: dict[str, dict], supplier: str) -> dict:
    """Find supplier entry, falling back to case-insensitive match."""
    if supplier in supplier_info:
        return supplier_info[supplier]
    sup_lower = supplier.strip().lower()
    for k, v in supplier_info.items():
        if k.strip().lower() == sup_lower:
            return v
    return {}


# ---------------------------------------------------------------------------
# Contract input xlsx
# ---------------------------------------------------------------------------

def _read_contract_input(xlsx_path: Path) -> tuple[list[str], list[list]]:
    """Return ``(headers, data_rows)`` from a contract input xlsx."""
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb.active
    headers = [_str(cell.value) for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    data = [list(row) for row in ws.iter_rows(min_row=2, values_only=True)]
    wb.close()
    return headers, data


# ---------------------------------------------------------------------------
# Word keyword replacement
# ---------------------------------------------------------------------------

def _clear_highlight(run) -> None:
    """Remove highlight and shading from *run* (strips highlighter-pen marks)."""
    from docx.oxml.ns import qn
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return
    for tag in (qn("w:highlight"), qn("w:shd")):
        el = rpr.find(tag)
        if el is not None:
            rpr.remove(el)


def _replace_paragraph(paragraph, replacements: dict[str, str], bold_keys: set[str] | None = None) -> None:
    """Replace keywords in *paragraph* across all runs (handles split runs)."""
    from docx.oxml.ns import qn

    # If the paragraph contains field characters (e.g. PAGE, NUMPAGES), use
    # safe run-by-run replacement to avoid corrupting the field XML structure.
    # Use iter() to search all descendants — fldChar lives inside w:r, not
    # directly under w:p, so find() would miss it.
    if next(paragraph._element.iter(qn("w:fldChar")), None) is not None:
        for run in paragraph.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)
                    _clear_highlight(run)
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    found_bold = False
    for key, value in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, value)
            if bold_keys and key in bold_keys:
                found_bold = True
    if new_text == full_text:
        return
    # Consolidate into the first run, wipe the rest
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
            run.font.name = "Arial"
            run.font.size = Pt(10)
            if found_bold:
                run.font.bold = True
            _clear_highlight(run)
        else:
            run.text = ""


def _iter_hf_containers(doc: Document, include_footers: bool = True):
    """Yield every header/footer container across all sections."""
    for section in doc.sections:
        headers = (section.header, section.first_page_header, section.even_page_header)
        footers = (section.footer, section.first_page_footer, section.even_page_footer)
        for hf in headers + (footers if include_footers else ()):
            try:
                yield hf
            except Exception:
                pass


def _replace_sdt_content(doc: Document, replacements: dict[str, str]) -> None:
    """Replace keywords inside Word Content Controls (w:sdt).

    Modifies w:t text directly within w:sdtContent and removes w:showingPlcHdr
    so Word displays the replaced content instead of the placeholder.
    """
    from docx.oxml.ns import qn

    xml_roots = [doc.element.body]
    for hf in _iter_hf_containers(doc):
        try:
            xml_roots.append(hf._element)
        except Exception:
            pass

    for root in xml_roots:
        for sdt in root.iter(qn("w:sdt")):
            sdt_content = sdt.find(qn("w:sdtContent"))
            if sdt_content is None:
                continue
            for t in sdt_content.iter(qn("w:t")):
                if not t.text:
                    continue
                for key, value in replacements.items():
                    if key in t.text:
                        t.text = t.text.replace(key, value)
            # Remove showingPlcHdr so Word shows content, not placeholder
            sdt_pr = sdt.find(qn("w:sdtPr"))
            if sdt_pr is not None:
                plc = sdt_pr.find(qn("w:showingPlcHdr"))
                if plc is not None:
                    sdt_pr.remove(plc)


def _replace_in_doc(doc: Document, replacements: dict[str, str], bold_keys: set[str] | None = None) -> None:
    """Replace keywords in body paragraphs, tables, and all section headers/footers."""

    def _process_paragraphs_and_tables(container) -> None:
        try:
            for para in container.paragraphs:
                _replace_paragraph(para, replacements, bold_keys)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_paragraph(para, replacements, bold_keys)
        except Exception:
            pass

    # Body
    _process_paragraphs_and_tables(doc)
    # Headers and footers
    for hf in _iter_hf_containers(doc):
        _process_paragraphs_and_tables(hf)


def _replace_xml_fallback(doc: Document, replacements: dict[str, str], bold_keys: set[str] | None = None) -> None:
    """Second-pass paragraph-level XML replacement.

    Iterates every ``w:p`` element across the document body and all
    header/footer XML roots, concatenates ALL nested ``w:t`` nodes within
    the paragraph (catches text inside ``w:hyperlink``, ``w:ins``, etc.),
    performs replacement, and writes the result back into the first ``w:t``
    while clearing the rest.  This handles keywords that are split across
    differently-structured runs that ``paragraph.runs`` would skip.
    """
    from docx.oxml.ns import qn

    xml_roots = [doc.element.body]
    for hf in _iter_hf_containers(doc):
        try:
            xml_roots.append(hf._element)
        except Exception:
            pass

    for root in xml_roots:
        for para_elem in root.iter(qn("w:p")):
            # Skip paragraphs that contain field characters
            if next(para_elem.iter(qn("w:fldChar")), None) is not None:
                continue
            # Only include w:t elements NOT inside a nested w:sdt
            # (SDT content is handled by _replace_sdt_content)
            t_elems = []
            for t in para_elem.iter(qn("w:t")):
                inside_sdt = False
                for anc in t.iterancestors():
                    if anc is para_elem:
                        break
                    if anc.tag == qn("w:sdt"):
                        inside_sdt = True
                        break
                if not inside_sdt:
                    t_elems.append(t)
            if not t_elems:
                continue
            full_text = "".join(t.text or "" for t in t_elems)
            new_text = full_text
            matched_bold = False
            for key, value in replacements.items():
                if key in new_text:
                    new_text = new_text.replace(key, value)
                    if bold_keys and key in bold_keys:
                        matched_bold = True
            if new_text == full_text:
                continue
            # Write result into first w:t, blank out the rest
            t_elems[0].text = new_text
            for t in t_elems[1:]:
                t.text = ""
            if matched_bold:
                r_elem = t_elems[0].getparent()
                if r_elem is not None and r_elem.tag == qn("w:r"):
                    rpr = r_elem.find(qn("w:rPr"))
                    if rpr is None:
                        rpr = r_elem.makeelement(qn("w:rPr"), {})
                        r_elem.insert(0, rpr)
                    if rpr.find(qn("w:b")) is None:
                        rpr.append(rpr.makeelement(qn("w:b"), {}))




def _find_product_table(doc: Document):
    """Locate the Product Rebate Table by header keywords."""
    for table in doc.tables:
        if not table.rows:
            continue
        joined = " ".join(_normalize(cell.text) for cell in table.rows[0].cells)
        if "segment" in joined and "rebate" in joined:
            return table
    return None


def _fill_row_cells(row, col_map: list[int | None], data_row: list, xlsx_headers: list[str] | None = None) -> None:
    """Write data values into a table row's cells (Times New Roman 8pt)."""
    for col_idx, cell in enumerate(row.cells):
        data_idx = col_map[col_idx] if col_idx < len(col_map) else None
        header = (xlsx_headers[data_idx] if xlsx_headers and data_idx is not None and data_idx < len(xlsx_headers) else "")
        text = _format_cell(data_row[data_idx], header) if (
            data_idx is not None and data_idx < len(data_row)
        ) else ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)


def _fill_product_table(table, xlsx_headers: list[str], data_rows: list[list]) -> None:
    """Fill the product table with *data_rows*.

    Rows 1+ are split into two groups:
    * **Keyword rows** (contain a ``<...>`` placeholder) — preserved; kept
      after the data so signature blocks are not deleted.
    * **Blank / template rows** (no placeholder) — deleted to make room.

    New data rows are inserted *before* the first preserved keyword row.
    """
    # Build column mapping: Word column index → xlsx column index
    word_headers = [_normalize(cell.text) for cell in table.rows[0].cells]
    xlsx_lower = [_normalize(h) for h in xlsx_headers]
    col_map: list[int | None] = []
    for wh in word_headers:
        match: int | None = None
        for i, dh in enumerate(xlsx_lower):
            if wh and (wh == dh or wh in dh or dh in wh):
                match = i
                break
        col_map.append(match)

    # Classify rows 1+: keep those with keyword placeholders, delete the rest
    tbl_elem = table._tbl
    keyword_trs: list = []
    for row in list(table.rows[1:]):
        row_text = " ".join(cell.text for cell in row.cells)
        if _KEYWORD_RE.search(row_text):
            keyword_trs.append(row._tr)
        else:
            tbl_elem.remove(row._tr)

    # Insert new data rows before the first keyword row (or append if none)
    import copy
    from docx.oxml.ns import qn as _qn3
    header_cells = table.rows[0].cells  # source for column width (w:tcPr)
    first_keyword_tr = keyword_trs[0] if keyword_trs else None
    for data_row in data_rows:
        new_row = table.add_row()
        # Copy w:tcPr from header row cells to preserve column widths
        for new_cell, hdr_cell in zip(new_row.cells, header_cells):
            hdr_tcPr = hdr_cell._tc.find(_qn3("w:tcPr"))
            if hdr_tcPr is not None:
                new_tc = new_cell._tc
                existing = new_tc.find(_qn3("w:tcPr"))
                if existing is not None:
                    new_tc.remove(existing)
                new_tc.insert(0, copy.deepcopy(hdr_tcPr))
        _fill_row_cells(new_row, col_map, data_row, xlsx_headers)
        if first_keyword_tr is not None:
            new_tr = new_row._tr
            tbl_elem.remove(new_tr)
            ins_idx = list(tbl_elem).index(first_keyword_tr)
            tbl_elem.insert(ins_idx, new_tr)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def _effective_date_from_fy_quarter(fy_2digit: int, quarter: int) -> str:
    """Return the first month label of the given HP FY quarter.

    Quarter mapping (HP fiscal year):
      Q1: Nov(fy-1), Dec(fy-1), Jan(fy)  → first month = Nov (fy-1)
      Q2: Feb, Mar, Apr                  → first month = Feb (fy)
      Q3: May, Jun, Jul                  → first month = May (fy)
      Q4: Aug, Sep, Oct                  → first month = Aug (fy)
    """
    full_year = 2000 + fy_2digit
    if quarter == 1:
        month, year = 11, full_year - 1
    elif quarter == 2:
        month, year = 2, full_year
    elif quarter == 3:
        month, year = 5, full_year
    else:  # Q4
        month, year = 8, full_year
    from datetime import datetime as _dt
    return _dt(year, month, 1).strftime("%B 1, %Y")


def generate_report(
    supplier_info_dir: Path,
    template_dir: Path,
    rebate_form_input_dir: Path,
    output_dir: Path,
    suppliers: list[str],
    form_numbers: dict[str, str],
    log: Callable[[str, str], None],
    fy: int | None = None,
    quarter: int | None = None,
    icertis_codes: dict[str, str] | None = None,
) -> tuple[list[Path], dict[str, list[str]]]:
    """Generate one Word contract per supplier.

    Returns a list of paths to the written ``.docx`` files.
    """
    # ── Supplier info Excel ───────────────────────────────────────────
    info_xlsx = _get_latest_file(supplier_info_dir, ".xlsx")
    if info_xlsx is None:
        log(f"No Excel found in supplier info folder: {supplier_info_dir}", "ERROR")
        return [], {}
    log(f"  Supplier info: {info_xlsx.name}", "INFO")
    supplier_info = _read_supplier_info(info_xlsx)

    # ── Word template ─────────────────────────────────────────────────
    template_docx = _get_latest_file(template_dir, ".docx")
    if template_docx is None:
        log(f"No Word template found in: {template_dir}", "ERROR")
        return [], {}
    log(f"  Template: {template_docx.name}", "INFO")

    # ── Output folder: <output_dir>/<YYYYMMDD>/ ───────────────────────
    today_str = date.today().strftime("%Y%m%d")
    out_dir = output_dir / today_str
    out_dir.mkdir(parents=True, exist_ok=True)

    out_paths: list[Path] = []
    blanks: dict[str, list[str]] = {}

    for supplier in suppliers:
        input_xlsx = rebate_form_input_dir / f"contract input - {supplier}.xlsx"
        if not input_xlsx.exists():
            log(f"  [{supplier}] contract input file not found — skipping", "WARNING")
            continue

        info = _lookup_supplier(supplier_info, supplier)
        if not info:
            log(f"  [{supplier}] not found in supplier info Excel — fields will be blank", "WARNING")

        def _val(key: str) -> str:
            """Case-insensitive column lookup from supplier info row."""
            if key in info:
                return _str(info[key])
            key_lower = key.lower()
            for k in info:
                if k.lower() == key_lower:
                    return _str(info[k])
            return ""

        replacements = {
            "ICMRebateAgreementCode": _val("ICMRebateAgreementCode"),
            "ICMAgreementCode":       (icertis_codes or {}).get(supplier, ""),
            "CW#":                    _val("CW#"),
            "ICMPartyName1":          _val("ICMPartyName1"),
            "ICMExternalSignatory":   _val("ICMExternalSignatory"),
            "ICMExternalSignatoryTitle": _val("ICMExternalSignatoryTitle"),
            "ICMAgreementNumber1":    form_numbers.get(supplier, ""),
            "ICMEffectiveDate":  (
                _effective_date_from_fy_quarter(fy, quarter)
                if fy is not None and quarter is not None
                else datetime.now().strftime("%b %Y")
            ),
        }

        xlsx_headers, xlsx_data = _read_contract_input(input_xlsx)

        # ── Detect blank cells in data columns ────────────────────────────────
        blank_cols: list[str] = []
        for col_idx, col_name in enumerate(xlsx_headers):
            if any(
                col_idx >= len(row) or row[col_idx] is None or str(row[col_idx]).strip() == ""
                for row in xlsx_data
            ):
                blank_cols.append(col_name)
        if blank_cols:
            blanks[supplier] = blank_cols
            log(f"  [{supplier}] Blank fields: {blank_cols}", "WARNING")

        doc = Document(template_docx)

        # Sort by key length descending to avoid substring replacement
        # (e.g. replace ICMExternalSignatoryTitle before ICMExternalSignatory)
        sorted_replacements = dict(
            sorted(replacements.items(), key=lambda kv: len(kv[0]), reverse=True)
        )

        # Warn about keywords that will be replaced with empty string
        missing_keys = [k for k, v in sorted_replacements.items() if not v]
        if missing_keys:
            log(f"  [{supplier}] Missing values for keywords: {missing_keys}", "WARNING")

        _replace_sdt_content(doc, sorted_replacements)   # first: convert SDTs to plain runs
        _replace_in_doc(doc, sorted_replacements)          # then: regular paragraph replacement
        _replace_xml_fallback(doc, sorted_replacements)    # finally: catch any remaining splits

        # Find first table anywhere in the document (including inside SDTs/containers)
        from docx.oxml.ns import qn as _qn2
        first_tbl = next(doc.element.body.iter(_qn2("w:tbl")), None)
        if first_tbl is None:
            log(f"  [{supplier}] No table found in template", "WARNING")
            product_table = None
        else:
            from docx.table import Table
            product_table = Table(first_tbl, doc)
        if product_table is not None:
            _fill_product_table(product_table, xlsx_headers, xlsx_data)

        fname = f"Rebate Agreement Update Form_{form_numbers.get(supplier, '')}_{supplier}.docx"
        out_path = out_dir / fname
        doc.save(out_path)
        log(f"  [{supplier}] Saved → {fname}", "INFO")
        out_paths.append(out_path)

    return out_paths, blanks
